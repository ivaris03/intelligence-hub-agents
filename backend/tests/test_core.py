import asyncio
from datetime import UTC, datetime, timedelta
from io import BytesIO

import pytest
from docx import Document
from PIL import Image
from pydantic import ValidationError
from sqlalchemy import select
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from app.agents.workflows import (
    DeepResearchHarness,
    Evidence,
    PresentationOutline,
    ResearchBudget,
    ResearchResult,
    ResearchSection,
    normalize_presentation_outline,
    presentation_content_page_limit,
    route_presentation_intent,
    validate_research_result,
)
from app.api.schemas import MessageRequest
from app.chat.service import should_search_web
from app.core.config import Settings
from app.core.security import contains_sensitive_memory, redact, remove_unverified_urls
from app.db.base import AppSettings, Base, Conversation, MemorySummary, Message
from app.files.service import FileValidationError, validate_upload
from app.integrations.qwen import QwenAdapter
from app.integrations.tavily import SearchResult, TavilyAdapter
from app.memory.service import parse_memory_command, refine_idle_memory_summary
from app.skills.service import normalize_skill_name


def test_mode_combinations_are_server_validated() -> None:
    with pytest.raises(ValidationError):
        MessageRequest(content="画图", mode="chat", agent_type="image")
    with pytest.raises(ValidationError):
        MessageRequest(content="画图", mode="work")
    assert MessageRequest(content="画图", mode="work", agent_type="image").agent_type == "image"


def test_thinking_effort_is_validated_and_mapped_to_budget() -> None:
    with pytest.raises(ValidationError):
        MessageRequest(content="分析", thinking_effort="extreme")

    settings = Settings(
        dashscope_api_key="test",
        tavily_api_key=None,
        qwen_thinking_budget=1024,
    )
    assert QwenAdapter(settings.with_thinking_effort("low")).thinking_parameters() == {
        "enable_thinking": True,
        "thinking_budget": 512,
    }
    assert QwenAdapter(settings.with_thinking_effort("medium")).thinking_parameters()[
        "thinking_budget"
    ] == 1024
    high_adapter = QwenAdapter(settings.with_thinking_effort("high"))
    assert high_adapter.thinking_parameters()["thinking_budget"] == 4096
    assert high_adapter.chat_model().extra_body == {
        "enable_thinking": True,
        "thinking_budget": 4096,
    }


def test_search_requires_explicit_language() -> None:
    assert should_search_web("请联网搜索今天的行业动态")
    assert should_search_web("search the web for current releases")
    assert not should_search_web("解释一下向量检索的原理")
    assert not should_search_web("这件事可能是最新方案")


def test_redaction_and_memory_safety() -> None:
    summary = redact({"Authorization": "Bearer secret-token", "query": "safe"})
    assert "secret-token" not in summary
    assert "[已脱敏]" in summary
    assert contains_sensitive_memory("我的 API_KEY 是 sk-example123456")
    assert parse_memory_command("你好，请记住我偏好简短回复") == (
        "remember",
        "我偏好简短回复",
    )


def test_upload_validation_checks_mime_and_magic() -> None:
    settings = Settings(dashscope_api_key=None, tavily_api_key=None)
    text = validate_upload("notes.md", "text/markdown", "你好".encode(), settings)
    assert text.kind == "document"
    with pytest.raises(FileValidationError):
        validate_upload("notes.md", "application/pdf", b"hello", settings)
    with pytest.raises(FileValidationError):
        validate_upload("fake.png", "image/png", b"not an image", settings)

    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("第二段")
    docx = BytesIO()
    document.save(docx)
    extracted = validate_upload(
        "brief.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        docx.getvalue(),
        settings,
    )
    assert [segment.locator for segment in extracted.segments] == ["第 1 段", "第 2 段"]

    png = BytesIO()
    Image.new("RGB", (32, 32), "green").save(png, format="PNG")
    assert validate_upload("reference.png", "image/png", png.getvalue(), settings).kind == "image"


def test_skill_normalization_and_slide_intent() -> None:
    assert normalize_skill_name("  Tech  Writing ") == "tech writing"
    assert route_presentation_intent("新建季度汇报") == "CREATE"
    assert route_presentation_intent("修改第 2 页", source_artifact_id="artifact") == "MODIFY"
    with pytest.raises(ValueError):
        route_presentation_intent("修改第 2 页")


def test_research_budget_and_citation_validation() -> None:
    budget = ResearchBudget(max_searches=1, timeout_seconds=30)
    asyncio.run(budget.consume_search())
    with pytest.raises(RuntimeError):
        asyncio.run(budget.consume_search())
    result = ResearchResult(
        title="测试",
        sections=[ResearchSection(heading="结论", content="内容", evidence_ids=["S1"])],
        evidence=[Evidence(id="S1", title="来源", url="https://example.com", snippet="摘要")],
    )
    assert validate_research_result(result) == []
    result.sections[0].evidence_ids = ["missing"]
    assert validate_research_result(result)


def test_provider_outputs_are_semantically_bounded() -> None:
    assert presentation_content_page_limit("制作三页以内的简报", 15) == 2
    outline = normalize_presentation_outline(
        PresentationOutline(title="title", slides=["title", "title", "验收结果"]),
        "P0 验收简报",
        2,
    )
    assert outline.title == "P0 验收简报"
    assert outline.slides == ["验收结果", "目标与范围"]

    checked, removed = remove_unverified_urls(
        "[官方](https://safe.example/docs) 与 https://fake.example/。",
        {"https://safe.example/docs"},
    )
    assert "https://safe.example/docs" in checked
    assert "https://fake.example" not in checked
    assert removed == ["https://fake.example/"]


def test_research_budget_exhaustion_is_a_tool_result_and_urls_are_grounded() -> None:
    async def scenario() -> None:
        settings = Settings(dashscope_api_key="test", tavily_api_key="test")
        harness = DeepResearchHarness(settings, ResearchBudget(max_searches=1, timeout_seconds=30))

        class FakeSearch:
            async def search(self, query: str, *, max_results: int = 5):
                del query, max_results
                return [SearchResult("官方", "https://example.com/docs", "可信摘要")]

        harness.adapter = FakeSearch()
        first = await harness.budgeted_search("测试")
        second = await harness.budgeted_search("再次测试")
        assert "https://example.com/docs" in first
        assert "search_budget_exhausted" in second
        grounded = harness._ground_result(
            ResearchResult(
                title="研究",
                sections=[
                    ResearchSection(
                        heading="结论",
                        content="可信内容 https://fake.example",
                        evidence_ids=["S1", "S2"],
                    )
                ],
                evidence=[
                    Evidence(
                        id="S1",
                        title="模型标题",
                        url="https://example.com/docs",
                        snippet="模型摘要",
                    ),
                    Evidence(
                        id="S2",
                        title="虚构",
                        url="https://fake.example",
                        snippet="虚构摘要",
                    ),
                ],
            ),
            "研究",
        )
        assert [item.url for item in grounded.evidence] == ["https://example.com/docs"]
        assert grounded.sections[0].evidence_ids == ["S1"]
        assert "https://fake.example" not in grounded.sections[0].content

    asyncio.run(scenario())


def test_qwen_embeddings_send_strings_and_validate_dimensions(monkeypatch) -> None:
    captured: dict[str, object] = {}

    class FakeEmbeddings:
        def __init__(self, **kwargs):
            captured.update(kwargs)

        async def aembed_documents(self, texts):
            captured["texts"] = texts
            return [[0.1, 0.2, 0.3]]

    monkeypatch.setattr("app.integrations.qwen.OpenAIEmbeddings", FakeEmbeddings)
    settings = Settings(
        dashscope_api_key="test",
        tavily_api_key=None,
        qwen_embedding_dimensions=3,
    )
    vectors = asyncio.run(QwenAdapter(settings).embed_documents(["原始字符串"]))
    assert captured["check_embedding_ctx_length"] is False
    assert captured["texts"] == ["原始字符串"]
    assert vectors == [[0.1, 0.2, 0.3]]


def test_qwen_stream_keeps_reasoning_separate_from_answer() -> None:
    assert QwenAdapter._stream_delta_parts(
        {
            "choices": [
                {"delta": {"reasoning_content": "先分析", "content": "再回答"}}
            ]
        }
    ) == [("reasoning", "先分析"), ("text", "再回答")]
    assert QwenAdapter._stream_delta_parts({"choices": []}) == []


def test_tavily_normalizes_nested_mcp_content() -> None:
    raw = {
        "content": [
            {
                "type": "text",
                "text": '{"results":[{"title":"官方文档","url":"https://example.com/docs","content":"摘要"}]}',
            }
        ]
    }
    assert TavilyAdapter._normalize(raw)[0].as_dict() == {
        "title": "官方文档",
        "url": "https://example.com/docs",
        "snippet": "摘要",
    }


def test_idle_memory_summary_refinement_is_cursor_based_safe_and_switchable() -> None:
    async def scenario() -> None:
        engine = create_async_engine("sqlite+aiosqlite://")
        sessions = async_sessionmaker(engine, expire_on_commit=False)
        async with engine.begin() as connection:
            await connection.run_sync(Base.metadata.create_all)
        now = datetime.now(UTC)
        async with sessions() as session:
            conversation = Conversation(last_activity_at=now - timedelta(minutes=31))
            session.add(conversation)
            await session.flush()
            session.add(
                Message(
                    conversation_id=conversation.id,
                    role="user",
                    content="我偏好简洁回答",
                    status="completed",
                    created_at=now - timedelta(minutes=32),
                )
            )
            await session.commit()
            assert await refine_idle_memory_summary(session, now=now) == 1
            assert await refine_idle_memory_summary(session, now=now) == 0
            saved = list((await session.scalars(select(MemorySummary))).all())
            assert len(saved) == 1
            assert saved[0].content == "我偏好简洁回答。"

            second = Conversation(last_activity_at=now - timedelta(minutes=31))
            session.add(second)
            await session.flush()
            session.add(
                Message(
                    conversation_id=second.id,
                    role="user",
                    content="我偏好详细回答",
                    status="completed",
                    created_at=now - timedelta(minutes=32),
                )
            )
            await session.commit()
            assert await refine_idle_memory_summary(session, now=now) == 0

            stored_settings = await session.get(AppSettings, 1)
            assert stored_settings is not None
            stored_settings.memory_enabled = False
            third = Conversation(last_activity_at=now - timedelta(minutes=31))
            session.add(third)
            await session.flush()
            session.add(
                Message(
                    conversation_id=third.id,
                    role="user",
                    content="我常用 Python",
                    status="completed",
                    created_at=now - timedelta(minutes=32),
                )
            )
            await session.commit()
            assert await refine_idle_memory_summary(session, now=now) == 0
        await engine.dispose()

    asyncio.run(scenario())
