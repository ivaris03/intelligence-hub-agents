from __future__ import annotations

import json
import math
import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from langchain_core.messages import HumanMessage
from PIL import Image, UnidentifiedImageError
from pydantic import BaseModel
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.config import Settings
from app.db.base import FileChunk, StoredFile
from app.files.storage import get_storage, make_storage_key
from app.integrations.qwen import QwenAdapter
from app.observability.langsmith import finish_trace, trace_operation


class FileValidationError(ValueError):
    pass


DOCUMENT_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
MIME_BY_EXTENSION = {
    ".txt": {"text/plain"},
    ".md": {"text/markdown", "text/plain", "text/x-markdown"},
    ".pdf": {"application/pdf"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    ".png": {"image/png"},
    ".jpg": {"image/jpeg"},
    ".jpeg": {"image/jpeg"},
    ".webp": {"image/webp"},
}


@dataclass(slots=True)
class ExtractedSegment:
    content: str
    locator: str


@dataclass(slots=True)
class ValidatedUpload:
    extension: str
    mime_type: str
    kind: str
    segments: list[ExtractedSegment]

    @property
    def text(self) -> str:
        return "\n\n".join(segment.content for segment in self.segments if segment.content)


class DocumentSelection(BaseModel):
    selected_indices: list[int]


def _validate_declared_mime(extension: str, declared_mime: str) -> str:
    mime = declared_mime.split(";", 1)[0].strip().lower()
    if mime not in MIME_BY_EXTENSION[extension]:
        raise FileValidationError("文件扩展名与 MIME 类型不一致")
    return mime


def _extract_document(extension: str, data: bytes) -> list[ExtractedSegment]:
    if extension in {".txt", ".md"}:
        if b"\x00" in data[:4096]:
            raise FileValidationError("文本文件包含不支持的二进制内容")
        try:
            text = data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise FileValidationError("文本文件必须使用 UTF-8 编码") from exc
        return [ExtractedSegment(text.strip(), f"字符 1-{len(text)}")]

    if extension == ".pdf":
        if not data.startswith(b"%PDF-"):
            raise FileValidationError("PDF 文件魔数无效")
        try:
            reader = PdfReader(BytesIO(data))
            return [
                ExtractedSegment((page.extract_text() or "").strip(), f"第 {index} 页")
                for index, page in enumerate(reader.pages, 1)
                if (page.extract_text() or "").strip()
            ]
        except Exception as exc:
            raise FileValidationError("PDF 文本解析失败") from exc

    if not data.startswith(b"PK"):
        raise FileValidationError("DOCX 文件魔数无效")
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            names = set(archive.namelist())
            if "word/document.xml" not in names or "[Content_Types].xml" not in names:
                raise FileValidationError("文件不是有效的 DOCX 文档")
        document = Document(BytesIO(data))
        segments = [
            ExtractedSegment(paragraph.text.strip(), f"第 {index} 段")
            for index, paragraph in enumerate(document.paragraphs, 1)
            if paragraph.text.strip()
        ]
        for table_index, table in enumerate(document.tables, 1):
            text = "\n".join(
                " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
            ).strip()
            if text:
                segments.append(ExtractedSegment(text, f"表格 {table_index}"))
        return segments
    except FileValidationError:
        raise
    except Exception as exc:
        raise FileValidationError("DOCX 文本解析失败") from exc


def validate_upload(
    filename: str, declared_mime: str, data: bytes, settings: Settings
) -> ValidatedUpload:
    extension = Path(filename).suffix.lower()
    if extension not in DOCUMENT_EXTENSIONS | IMAGE_EXTENSIONS:
        raise FileValidationError("不支持该文件格式")
    if not data:
        raise FileValidationError("文件内容为空")
    if len(data) > settings.max_upload_bytes:
        raise FileValidationError("单个文件不能超过 20 MB")
    mime = _validate_declared_mime(extension, declared_mime)

    if extension in DOCUMENT_EXTENSIONS:
        segments = _extract_document(extension, data)
        if not any(segment.content for segment in segments):
            raise FileValidationError("文档中没有可提取的文本")
        return ValidatedUpload(extension, mime, "document", segments)

    try:
        Image.MAX_IMAGE_PIXELS = settings.max_image_pixels
        with Image.open(BytesIO(data)) as image:
            image.verify()
        with Image.open(BytesIO(data)) as image:
            width, height = image.size
            image_format = (image.format or "").lower()
    except (UnidentifiedImageError, OSError, Image.DecompressionBombError) as exc:
        raise FileValidationError("图片内容无法解码或尺寸不安全") from exc
    if width * height > settings.max_image_pixels:
        raise FileValidationError("图片像素尺寸超过限制")
    expected = "jpeg" if extension in {".jpg", ".jpeg"} else extension[1:]
    if image_format != expected:
        raise FileValidationError("图片扩展名与文件魔数不一致")
    return ValidatedUpload(extension, mime, "image", [])


def _split_segment(segment: ExtractedSegment, size: int, overlap: int) -> list[ExtractedSegment]:
    text = segment.content.strip()
    if len(text) <= size:
        return [segment]
    pieces: list[ExtractedSegment] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            boundary = max(text.rfind("\n", start, end), text.rfind("。", start, end))
            if boundary > start + size // 2:
                end = boundary + 1
        pieces.append(
            ExtractedSegment(text[start:end].strip(), f"{segment.locator} · 字符 {start + 1}-{end}")
        )
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return pieces


def document_chunks_for_upload(
    upload: ValidatedUpload, settings: Settings
) -> list[ExtractedSegment]:
    if upload.kind != "document":
        return []
    return [
        piece
        for segment in upload.segments
        for piece in _split_segment(
            segment, settings.document_chunk_chars, settings.document_chunk_overlap
        )
        if piece.content
    ]


async def create_file(
    session: AsyncSession,
    settings: Settings,
    conversation_id: UUID,
    filename: str,
    declared_mime: str,
    data: bytes,
) -> StoredFile:
    validated = validate_upload(filename, declared_mime, data, settings)
    key = make_storage_key("uploads", validated.extension)
    stored = StoredFile(
        conversation_id=conversation_id,
        name=Path(filename).name[:255],
        mime_type=validated.mime_type,
        kind=validated.kind,
        size=len(data),
        storage_key=key,
        status="processing",
        text_content=(validated.text if validated.kind == "document" else None),
    )
    session.add(stored)
    await session.flush()
    storage = get_storage(settings)
    try:
        await storage.save(key, data, validated.mime_type)
        if validated.kind == "document":
            chunks = document_chunks_for_upload(validated, settings)
            embeddings: list[list[float] | None]
            try:
                embeddings = await QwenAdapter(settings).embed_documents(
                    [chunk.content for chunk in chunks]
                )
            except Exception:
                embeddings = [None] * len(chunks)
            session.add_all(
                FileChunk(
                    file_id=stored.id,
                    chunk_index=index,
                    content=chunk.content,
                    locator=chunk.locator,
                    embedding=embeddings[index],
                )
                for index, chunk in enumerate(chunks)
            )
        stored.status = "ready"
        await session.commit()
    except Exception:
        stored.status = "failed"
        stored.error = "文件处理失败，请重试。"
        await session.commit()
        await storage.delete(key)
        raise
    await session.refresh(stored)
    return stored


async def load_files_for_request(
    session: AsyncSession,
    conversation_id: UUID,
    file_ids: list[UUID],
    settings: Settings,
) -> list[StoredFile]:
    if len(file_ids) > settings.max_files_per_request:
        raise FileValidationError("单条请求最多关联 3 个文件")
    if not file_ids:
        return []
    result = await session.scalars(
        select(StoredFile)
        .where(
            StoredFile.id.in_(file_ids),
            StoredFile.conversation_id == conversation_id,
            StoredFile.status == "ready",
        )
        .options(selectinload(StoredFile.chunks))
    )
    by_id = {file.id: file for file in result.unique().all()}
    if any(file_id not in by_id for file_id in file_ids):
        raise FileValidationError("文件不存在、尚未就绪或不属于当前会话")
    return [by_id[file_id] for file_id in file_ids]


def _terms(text: str) -> set[str]:
    lower = text.lower()
    words = set(re.findall(r"[a-z0-9_]{2,}|[\u4e00-\u9fff]", lower))
    chinese = "".join(re.findall(r"[\u4e00-\u9fff]", lower))
    words.update(chinese[index : index + 2] for index in range(max(0, len(chinese) - 1)))
    return words


def _cosine(left: list[float], right: list[float]) -> float:
    numerator = sum(a * b for a, b in zip(left, right, strict=False))
    denominator = math.sqrt(sum(a * a for a in left)) * math.sqrt(sum(b * b for b in right))
    return numerator / denominator if denominator else 0.0


async def rerank_document_items(
    query: str,
    items: list[dict[str, Any]],
    settings: Settings,
    *,
    limit: int,
) -> list[dict[str, Any]]:
    """Rerank a bounded candidate set by direct answerability with a stable fallback."""

    if len(items) <= limit or not settings.model_ready:
        return items[:limit]
    payload = [
        {"index": index, "text": str(item.get("text") or "")[:2_000]}
        for index, item in enumerate(items)
    ]
    try:
        selector = QwenAdapter(settings).chat_model(work=True).with_structured_output(
            DocumentSelection
        )
        selected = DocumentSelection.model_validate(
            await selector.ainvoke(
                [
                    HumanMessage(
                        content=(
                            "按文本能否直接、完整回答精确问题来排序候选文档。"
                            "不要因为共享‘支持、上传、格式’等泛化词就优先选择相邻主题；"
                            f"只返回最相关的 {limit} 个零基 index，不重复。\n"
                            f"问题：{query}\n候选：{json.dumps(payload, ensure_ascii=False)}"
                        )
                    )
                ],
                config={
                    "run_name": "document.semantic_rerank",
                    "tags": ["rag", "rerank"],
                },
            )
        ).selected_indices
    except Exception:
        return items[:limit]
    ordered: list[dict[str, Any]] = []
    seen: set[int] = set()
    for index in [*selected, *range(len(items))]:
        if index in seen or not 0 <= index < len(items):
            continue
        seen.add(index)
        ordered.append(items[index])
        if len(ordered) == limit:
            break
    return ordered


async def document_context(
    session: AsyncSession,
    files: list[StoredFile],
    query: str,
    settings: Settings,
    limit: int = 6,
) -> tuple[str, list[dict[str, Any]]]:
    with trace_operation(
        settings,
        "document.retrieve",
        inputs={
            "query": query,
            "limit": limit,
            "files": [{"id": str(file.id), "name": file.name, "kind": file.kind} for file in files],
        },
        run_type="retriever",
        tags=["rag", "retrieval"],
        metadata={"embedding_model": settings.qwen_embedding_model, "top_k": limit},
    ) as trace:
        context, sources = await _document_context(session, files, query, settings, limit)
        documents: list[dict[str, Any]] = []
        for source in sources:
            file = next((item for item in files if str(item.id) == source["file_id"]), None)
            if file is None:
                continue
            chunk = next(
                (item for item in file.chunks if item.locator == source["locator"]),
                None,
            )
            page_content = chunk.content if chunk else file.text_content or ""
            documents.append(
                {
                    "page_content": page_content[:8_000],
                    "metadata": source,
                }
            )
        finish_trace(trace, {"documents": documents})
        return context, sources


async def _document_context(
    session: AsyncSession,
    files: list[StoredFile],
    query: str,
    settings: Settings,
    limit: int = 6,
) -> tuple[str, list[dict[str, Any]]]:
    documents = [file for file in files if file.kind == "document"]
    if not documents:
        return "", []
    inline = [
        file
        for file in documents
        if file.text_content and len(file.text_content) <= settings.document_inline_chars
    ]
    sources: list[dict[str, Any]] = []
    context_blocks: list[str] = []
    for file in inline:
        locator = file.chunks[0].locator if file.chunks else "全文"
        context_blocks.append(f"[文件：{file.name}；{locator}]\n{file.text_content}")
        sources.append({"file_id": str(file.id), "name": file.name, "locator": locator})

    long_files = [file for file in documents if file not in inline]
    long_chunks = [chunk for file in long_files for chunk in file.chunks]
    if long_chunks:
        query_embedding: list[float] | None = None
        try:
            embedded = await QwenAdapter(settings).embed_documents([query])
            query_embedding = embedded[0]
        except Exception:
            pass
        query_terms = _terms(query)

        def score(chunk: FileChunk) -> float:
            lexical = len(query_terms & _terms(chunk.content)) / max(1, len(query_terms))
            vector = (
                _cosine(query_embedding, list(chunk.embedding))
                if query_embedding is not None and chunk.embedding is not None
                else 0.0
            )
            return vector * 0.7 + lexical * 0.3

        ranked: list[FileChunk]
        dialect = session.bind.dialect.name if session.bind is not None else ""
        if query_embedding is not None and dialect == "postgresql":
            # Keep vector similarity inside PostgreSQL for real deployments. The
            # lexical fallback below keeps the no-key SQLite demo deterministic.
            vector_candidates = list(
                (
                    await session.scalars(
                        select(FileChunk)
                        .where(
                            FileChunk.file_id.in_([file.id for file in long_files]),
                            FileChunk.embedding.is_not(None),
                        )
                        .order_by(FileChunk.embedding.cosine_distance(query_embedding))
                        .limit(max(limit * 4, limit))
                    )
                ).all()
            )
            lexical_candidates = sorted(long_chunks, key=score, reverse=True)[: limit * 4]
            candidates = {chunk.id: chunk for chunk in vector_candidates}
            candidates.update({chunk.id: chunk for chunk in lexical_candidates})
            ranked = sorted(candidates.values(), key=score, reverse=True)[: limit * 4]
        else:
            ranked = sorted(long_chunks, key=score, reverse=True)[: limit * 4]
        reranked = await rerank_document_items(
            query,
            [{"text": chunk.content, "chunk": chunk} for chunk in ranked],
            settings,
            limit=limit,
        )
        ranked = [item["chunk"] for item in reranked]
        file_by_id = {file.id: file for file in documents}
        for chunk in ranked:
            file = file_by_id[chunk.file_id]
            context_blocks.append(f"[文件：{file.name}；{chunk.locator}]\n{chunk.content}")
            sources.append({"file_id": str(file.id), "name": file.name, "locator": chunk.locator})
    return "\n\n".join(context_blocks), sources


async def image_inputs(files: list[StoredFile], settings: Settings) -> list[dict[str, str]]:
    storage = get_storage(settings)
    inputs: list[dict[str, str]] = []
    for file in files:
        if file.kind != "image":
            continue
        data = await storage.read(file.storage_key)
        import base64

        encoded = base64.b64encode(data).decode("ascii")
        inputs.append(
            {
                "file_id": str(file.id),
                "name": file.name,
                "data_url": f"data:{file.mime_type};base64,{encoded}",
            }
        )
    return inputs
